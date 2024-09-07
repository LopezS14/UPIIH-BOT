import random
import io
import streamlit as st
import json
import pickle
import numpy as np
import pandas as pd
from datetime import datetime
from docx import Document as DocxDocument
from keras.models import load_model
import nltk
from nltk.stem import WordNetLemmatizer

lemmatizer = WordNetLemmatizer()

# Importamos los archivos generados en el código anterior
intents = json.loads(open('intents.json').read())
words = pickle.load(open('words.pkl', 'rb'))
classes = pickle.load(open('classes.pkl', 'rb'))
model = load_model('chatbot_model.h5')

# Pasamos las palabras de oración a su forma raíz
def clean_up_sentence(sentence):
    sentence_words = nltk.word_tokenize(sentence)
    sentence_words = [lemmatizer.lemmatize(word) for word in sentence_words]
    return sentence_words

# Convertimos la información a unos y ceros según si están presentes en los patrones
def bag_of_words(sentence):
    sentence_words = clean_up_sentence(sentence)
    bag = [0]*len(words)
    for w in sentence_words:
        for i, word in enumerate(words):
            if word == w:
                bag[i] = 1
    return np.array(bag)

# Predecimos la categoría a la que pertenece la oración
def predict_class(sentence):
    bow = bag_of_words(sentence)
    res = model.predict(np.array([bow]))[0]
    max_index = np.where(res == np.max(res))[0][0]
    category = classes[max_index]
    return category

# Obtenemos una respuesta aleatoria
def get_response(tag, intents_json):
    list_of_intents = intents_json['intents']
    result = ""
    for i in list_of_intents:
        if i["tag"] == tag:
            result = random.choice(i['responses'])
            break

    # Si el tag es "Mecatronica semestre 1", manejar el documento y proporcionar el botón de descarga
    if tag == "Mecatronica semestre  1":
        # st.write('Aquí tienes, si necesitas algo más dímelo')
        result = random.choice(i['responses'])

        # Cargar el documento existente
        doc_path = "C:/Users/Deyanira LS/Desktop/Bot/semestre7.docx"
        doc = DocxDocument(doc_path)
        
        # Acceder a la tabla específica
        if len(doc.tables) > 0:
            tabla_fecha = doc.tables[0]  # Acceder a la primera tabla
            
            if len(tabla_fecha.rows) > 0 and len(tabla_fecha.columns) > 1:
                fila_index = 0  # Índice de la fila (comenzando desde 0)
                columna_index = 1  # Índice de la columna (comenzando desde 0)
                
                # Accede a la celda específica
                celda = tabla_fecha.rows[fila_index].cells[columna_index]
                current_date = datetime.today().strftime('%m/%d/%y %H:%M:%S')
                celda.text = f"Fecha: {current_date}"

        # Guardar en un buffer en memoria
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)  # Regresa al inicio del buffer
        
        
         # Si el tag es "Mecatronica semestre 2 ", manejar el documento y proporcionar el botón de descarga
    if tag == "Mecatronica semestre  2":
        # st.write('Aquí tienes, si necesitas algo más dímelo')
        result = random.choice(i['responses'])

        # Cargar el documento existente
        Mecatronica_2 = "C:/Users/Deyanira LS/Desktop/Bot/semestre7.docx"
        doc = DocxDocument(Mecatronica_2)
        
        # Acceder a la tabla específica
        if len(doc.tables) > 0:
            tabla_fecha = doc.tables[0]  # Acceder a la primera tabla
            
            if len(tabla_fecha.rows) > 0 and len(tabla_fecha.columns) > 1:
                fila_index = 0  # Índice de la fila (comenzando desde 0)
                columna_index = 1  # Índice de la columna (comenzando desde 0)
                
                # Accede a la celda específica
                celda = tabla_fecha.rows[fila_index].cells[columna_index]
                current_date = datetime.today().strftime('%m/%d/%y %H:%M:%S')
                celda.text = f"Fecha: {current_date}"

        # Guardar en un buffer en memoria
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)  # Regresa al inicio del buffer
        
        # Mostrar el botón de descarga en Streamlit
        st.download_button(
            label="Descargar Temario",
            data=buffer,
            file_name="Temario_Actualizado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    return result
