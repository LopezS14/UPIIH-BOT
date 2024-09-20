import random
import io
import streamlit as st
import json
import pickle
import numpy as np
import requests
from datetime import datetime
from docx import Document as DocxDocument
import re
from tensorflow.keras.models import load_model

# Función para descargar archivos desde URLs
def download_file(url, local_filename):
    response = requests.get(url)
    response.raise_for_status()
    with open(local_filename, 'wb') as f:
        f.write(response.content)

# URLs de los archivos
intents_url = 'https://raw.githubusercontent.com/LopezS14/UPIIH-BOT/6589c20626f5d7425d3dbdfd2f8de8e1a6f81ab5/Bot/intents.json'
words_url = 'https://raw.githubusercontent.com/LopezS14/UPIIH-BOT/641fb9f82b895e8d8492c25abd0e15a82a5be90f/words.pkl'
classes_url = 'https://raw.githubusercontent.com/LopezS14/UPIIH-BOT/641fb9f82b895e8d8492c25abd0e15a82a5be90f/classes.pkl'
model_url = 'https://raw.githubusercontent.com/LopezS14/UPIIH-BOT/641fb9f82b895e8d8492c25abd0e15a82a5be90f/chatbot_model.h5'

# Descargar archivos
download_file(intents_url, 'intents.json')
download_file(words_url, 'words.pkl')
download_file(classes_url, 'classes.pkl')
download_file(model_url, 'chatbot_model.h5')

# Cargar archivos locales
with open('intents.json') as f:
    intents = json.load(f)
words = pickle.load(open('words.pkl', 'rb'))
classes = pickle.load(open('classes.pkl', 'rb'))
model = load_model('chatbot_model.h5')  # Cargar el modelo de la ruta

# Función que tokeniza usando expresiones regulares en lugar de NLTK
def clean_up_sentence(sentence):
    sentence_words = re.findall(r'\b\w+\b', sentence.lower())
    return sentence_words

# Convertimos la información a unos y ceros según si están presentes en los patrones
def bag_of_words(sentence):
    sentence_words = clean_up_sentence(sentence)
    bag = [0] * len(words)
    for w in sentence_words:
        for i, word in enumerate(words):
            if word == w:
                bag[i] = 1
    return np.array(bag)

# Predecimos la categoría a la que pertenece la oración
def predict_class(sentence):
    bow = bag_of_words(sentence)
    res = model.predict(np.array([bow]))[0]
    max_index = np.where(res == np.max(res))[0][0]
    category = classes[max_index]
    return category

# Diccionario de rutas de documentos
doc_paths = {
    
    "Ingenieria mecatronica semestre 1-programasintetico": "https://raw.githubusercontent.com/LopezS14/UPIIH-BOT/c80bb2b0d43be9b63d1bd0f2fc14f2458af6266c/Bot/M_PS1.pdf",
    "Ingenieria mecatronica semestre 1": "https://raw.githubusercontent.com/LopezS14/UPIIH-BOT/c80bb2b0d43be9b63d1bd0f2fc14f2458af6266c/Bot/ingenieriaMecatronica_1.docx"
}


# Función para manejar el documento y proporcionar el botón de descarga
def handle_document(tag):
    result = ""
    doc_path = doc_paths.get(tag)

    if not doc_path:
        return "Documento no encontrado para el semestre o carrera solicitado."

    st.write(f"Procesando documento: {doc_path}")

    try:
        if doc_path.endswith('.docx'):
            doc = DocxDocument(doc_path)
            if len(doc.tables) > 0:
                tabla_fecha = doc.tables[0]
                if len(tabla_fecha.rows) > 0 and len(tabla_fecha.columns) > 1:
                    fila_index = 0
                    columna_index = 1
                    celda = tabla_fecha.rows[fila_index].cells[columna_index]
                    current_date = datetime.today().strftime('%m/%d/%y %H:%M:%S')
                    celda.text = f"Fecha: {current_date}"

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="Descargar Temario",
                data=buffer,
                file_name=f"Temario_{tag.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        elif doc_path.endswith('.pdf'):
            with open(doc_path, "rb") as pdf_file:
                pdf_buffer = pdf_file.read()

            st.download_button(
                label="Descargar programa sintetico",
                data=pdf_buffer,
                file_name=f"Programa_Sintetico_{tag.replace(' ', '_')}.pdf",
                mime="application/pdf"
            )

    except Exception as e:
        result = f"Ocurrió un error al manejar el archivo: {e}"
        st.error(result)
    return result

# Obtener una respuesta aleatoria
def get_response(tag, intents_json):
    list_of_intents = intents_json['intents']
    result = ""
    for i in list_of_intents:
        if i["tag"] == tag:
            result = random.choice(i['responses'])
            break

    if tag in doc_paths:
        doc_result = handle_document(tag)
        result += "\n" + doc_result  # Agregar resultado del manejo de documento

    return result if result else "No se encontró una respuesta adecuada."

# Ejemplo de uso
# st.write(get_response("Sistemas automotrices semestre 7", intents))
